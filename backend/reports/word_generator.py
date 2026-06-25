from pathlib import Path

from docx import Document


def generate_word_report(
    detections
):

    reports_dir = Path(
        "generated_reports"
    )

    reports_dir.mkdir(
        exist_ok=True
    )

    report_path = (

        reports_dir /

        "forensic_report.docx"
    )

    doc = Document()

    doc.add_heading(

        "A3-ENDS Forensic Report",

        level=1

    )

    doc.add_paragraph(

        f"Total Incidents: "
        f"{len(detections)}"
    )

    for detection in detections:

        doc.add_paragraph(

            str(detection)
        )

    doc.save(
        report_path
    )

    return str(
        report_path
    )